from __future__ import annotations

from pathlib import Path

from docx import Document


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def add_bullets(doc: Document, items: list[str]):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def upgrade_report(src_path: Path, out_path: Path) -> None:
    doc = Document(str(src_path))

    # Fix common cover typos if present
    for p in doc.paragraphs[:40]:
        if "Danh s" in p.text and "th" in p.text:
            p.text = "Danh sách thành viên:"
        if p.text.strip() == "Giảng viên hướng dẫn :":
            p.text = "Giảng viên hướng dẫn:"

    # Find the first "Giới thiệu" heading and remove everything after it.
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Giới thiệu":
            start_idx = i
            break

    if start_idx is None:
        raise RuntimeError('Could not find heading "Giới thiệu" in source document')

    # Delete from start_idx to end
    for p in list(doc.paragraphs[start_idx:]):
        delete_paragraph(p)

    # Rebuild content following the same frame, but expanded.
    doc.add_heading("Giới thiệu", level=1)

    doc.add_heading("Đặt vấn đề", level=2)
    doc.add_paragraph(
        "Trong bối cảnh chuyển đổi số, người dùng ngày càng ưu tiên tìm kiếm, so sánh và đặt tour du lịch trực tuyến. "
        "Các thao tác trước đây thường thực hiện qua điện thoại hoặc tại quầy (hỏi tư vấn, xác nhận lịch khởi hành, thanh toán) "
        "nay được kỳ vọng có thể xử lý nhanh gọn trên website, đồng thời cung cấp thông tin minh bạch và dễ tra cứu."
    )
    doc.add_paragraph(
        "Tuy nhiên, thực tế cho thấy nhiều website du lịch chỉ tập trung hiển thị danh sách sản phẩm, chưa tối ưu luồng đặt tour "
        "theo lịch khởi hành và điểm đón, cũng như thiếu công cụ quản trị tập trung cho doanh nghiệp (tour, lịch trình, đơn đặt, thanh toán). "
        "Vì vậy, đồ án hướng tới xây dựng một hệ thống booking tour có cấu trúc rõ ràng, dễ mở rộng và đủ nghiệp vụ để mô phỏng triển khai thực tế."
    )

    doc.add_heading("Các giải pháp đã có", level=2)
    doc.add_paragraph(
        "Thị trường có nhiều nền tảng đặt dịch vụ du lịch như Booking.com, Agoda, Traveloka, Klook hoặc các website riêng của công ty lữ hành. "
        "Các nền tảng này thường cung cấp chức năng tìm kiếm, xem chi tiết, đặt dịch vụ và thanh toán trực tuyến."
    )
    doc.add_paragraph("Một số hạn chế thường gặp:")
    add_bullets(
        doc,
        [
            "Giao diện và luồng thao tác có thể phức tạp với người dùng phổ thông; nhiều bước dẫn tới giảm tỷ lệ chuyển đổi.",
            "Thông tin tour đôi khi trình bày rời rạc, khó so sánh giữa các lựa chọn (giá, lịch khởi hành, lịch trình, dịch vụ).",
            "Doanh nghiệp vừa và nhỏ chịu phí nền tảng/hoa hồng cao và bị giới hạn khả năng tùy biến theo nghiệp vụ riêng.",
            "Một số hệ thống chưa hỗ trợ đầy đủ vòng đời đơn đặt: thanh toán, cập nhật trạng thái, check-in, đánh giá sau chuyến đi.",
        ],
    )

    doc.add_heading("Giải pháp đề xuất", level=2)
    doc.add_paragraph(
        "ZakiBooking là website booking tour du lịch trong nước và quốc tế phát triển bằng Java Spring Boot (Spring MVC, Spring Data JPA, "
        "Spring Security) kết hợp Thymeleaf để render giao diện phía server. Hệ thống được tổ chức theo kiến trúc MVC phân tầng "
        "(Controller - Service - Repository - Entity) nhằm đảm bảo tính rõ ràng, dễ bảo trì và dễ mở rộng."
    )
    doc.add_paragraph("Các điểm nổi bật của giải pháp:")
    add_bullets(
        doc,
        [
            "Luồng đặt tour hoàn chỉnh: xem tour → chọn ngày khởi hành/điểm đón → áp mã giảm giá → tạo booking → thanh toán VNPay (sandbox).",
            "Tự động cập nhật trạng thái thanh toán qua return/callback (IPN) thay vì cập nhật thủ công.",
            "QR check-in cho mỗi booking, hỗ trợ xác thực nhanh khi bắt đầu chuyến đi.",
            "Đăng nhập truyền thống và OAuth2 (Google/Facebook) giúp tăng trải nghiệm đăng nhập.",
            "Tích hợp API bên thứ ba: Amadeus (hỗ trợ dữ liệu vé/chuyến bay), News API (tin tức), AI Chatbot (tư vấn theo dữ liệu hệ thống).",
            "Phân hệ Admin có dashboard thống kê và các màn hình CRUD quản trị tour, ngày khởi hành, mã giảm giá, người dùng, booking, đánh giá, liên hệ.",
        ],
    )

    doc.add_heading("Thiết kế và triển khai", level=1)

    doc.add_heading("Các yêu cầu chức năng", level=2)
    doc.add_paragraph(
        "Hệ thống phục vụ hai nhóm người dùng: (1) Khách hàng và (2) Quản trị viên. Dưới đây là các yêu cầu chức năng chính đã triển khai trong dự án."
    )

    doc.add_paragraph("2.1.1. Quản lý người dùng và xác thực")
    add_bullets(
        doc,
        [
            "Đăng ký, đăng nhập, đăng xuất; quản lý phiên đăng nhập.",
            "Mã hóa mật khẩu bằng BCrypt.",
            "Đăng nhập OAuth2 (Google/Facebook) khi cấu hình thông tin client.",
            "Phân quyền USER/ADMIN; giới hạn truy cập trang quản trị.",
        ],
    )

    doc.add_paragraph("2.1.2. Tìm kiếm và xem thông tin tour")
    add_bullets(
        doc,
        [
            "Trang danh sách tour hiển thị thông tin tóm tắt: tên tour, giá, hình ảnh, đánh giá, số lượt đặt (nếu có).",
            "Hỗ trợ lọc/tìm kiếm theo tiêu chí phổ biến: điểm đi/đến, ngày khởi hành, khoảng giá, phương tiện (tùy dữ liệu cấu hình).",
            "Trang chi tiết tour hiển thị mô tả, lịch trình theo ngày, danh sách ngày khởi hành và các thông tin liên quan.",
        ],
    )

    doc.add_paragraph("2.1.3. Đặt tour và quản lý đơn đặt")
    add_bullets(
        doc,
        [
            "Người dùng chọn ngày khởi hành, điểm đón, số lượng hành khách và nhập thông tin liên hệ.",
            "Áp mã giảm giá (voucher) nếu hợp lệ.",
            "Tạo đơn đặt (booking) và lưu vào MySQL.",
            "Người dùng theo dõi lịch sử booking, trạng thái thanh toán và thông tin check-in.",
        ],
    )

    doc.add_paragraph("2.1.4. Thanh toán VNPay")
    add_bullets(
        doc,
        [
            "Tạo URL thanh toán VNPay sandbox.",
            "Nhận kết quả từ VNPay (return) và cập nhật trạng thái booking.",
            "Xử lý callback/IPN để đảm bảo trạng thái thanh toán được cập nhật tin cậy.",
        ],
    )

    doc.add_paragraph("2.1.5. Yêu thích, đánh giá và tương tác")
    add_bullets(
        doc,
        [
            "Yêu thích (Favorites): thêm/bỏ tour khỏi danh sách yêu thích.",
            "Đánh giá (Reviews): người dùng gửi số sao và nội dung nhận xét.",
            "Liên hệ: gửi form liên hệ/yêu cầu tư vấn.",
        ],
    )

    doc.add_paragraph("2.1.6. Quản trị hệ thống (Admin)")
    add_bullets(
        doc,
        [
            "Dashboard: thống kê tổng quan doanh thu/đơn đặt/hiệu suất tour.",
            "Quản lý tour: CRUD tour, trạng thái, ảnh, thông tin mở rộng.",
            "Quản lý lịch trình: CRUD lịch trình theo tour.",
            "Quản lý ngày khởi hành: thêm/sửa ngày khởi hành và cấu hình điểm đón × ngày khởi hành.",
            "Quản lý mã giảm giá: tạo/cập nhật điều kiện áp dụng.",
            "Quản lý booking: xem danh sách, xem chi tiết, theo dõi trạng thái thanh toán.",
            "Quản lý đánh giá và liên hệ: xem/kiểm duyệt nội dung phản hồi.",
        ],
    )

    doc.add_heading("Các yêu cầu phi chức năng", level=2)
    add_bullets(
        doc,
        [
            "Tính dễ sử dụng: giao diện rõ ràng, responsive; giảm số bước trong luồng đặt tour.",
            "Hiệu năng: tối ưu số truy vấn, đảm bảo tải trang danh sách/chi tiết ổn định.",
            "Bảo mật: xác thực và phân quyền bằng Spring Security; mã hóa mật khẩu; hạn chế truy cập admin.",
            "Khả năng mở rộng: tổ chức theo tầng giúp dễ bổ sung thanh toán, báo cáo, thông báo.",
            "Độ tin cậy: dùng MySQL và Flyway để đảm bảo dữ liệu nhất quán và schema đồng bộ.",
        ],
    )

    doc.add_heading("Các ràng buộc", level=2)
    add_bullets(
        doc,
        [
            "Công nghệ: Java 17, Spring Boot 3.x, Thymeleaf, MySQL 8.x.",
            "Thời gian: triển khai theo phạm vi đồ án, ưu tiên luồng nghiệp vụ chính.",
            "Tích hợp: VNPay/Amadeus/News/AI phụ thuộc cấu hình API key và môi trường mạng.",
        ],
    )

    doc.add_heading("Mô hình hệ thống / Thiết kế giải pháp", level=2)
    doc.add_paragraph(
        "Hệ thống áp dụng kiến trúc MVC phân tầng. Controller chịu trách nhiệm điều hướng view và cung cấp API cho UI; Service xử lý nghiệp vụ; "
        "Repository thao tác dữ liệu qua JPA; Entity ánh xạ dữ liệu trong MySQL. Các tích hợp bên thứ ba (VNPay, Amadeus, News, AI) được tách riêng để giảm phụ thuộc."
    )
    doc.add_paragraph("Các thực thể (Entity) quan trọng trong dự án:")
    add_bullets(
        doc,
        [
            "NguoiDung: thông tin tài khoản, vai trò (USER/ADMIN), provider OAuth.",
            "ChuyenDi: thông tin tour (tiêu đề, mô tả, giá, trạng thái, liên kết điểm đến...).",
            "LichTrinh: lịch trình chi tiết theo ngày của tour.",
            "NgayKhoiHanh: danh sách ngày khởi hành theo tour.",
            "DiemDon: điểm đón khách.",
            "NgayKhoiHanhDiemDon: dữ liệu cấu hình/báo giá theo điểm đón × ngày khởi hành.",
            "DatCho: booking (đơn đặt), liên kết người dùng - tour - ngày khởi hành, trạng thái thanh toán.",
            "MaGiamGia, DanhGia, YeuThich, Contact: các thành phần tương tác và khuyến mãi.",
        ],
    )
    doc.add_paragraph(
        "Dữ liệu được quản lý bằng Flyway migration giúp dễ triển khai trên local và môi trường cloud (ví dụ Railway) mà không cần thao tác import thủ công."
    )

    doc.add_heading("Một số thành phần khác của đồ án", level=1)

    doc.add_heading("Kế hoạch dự án", level=2)
    doc.add_paragraph("Kế hoạch thực hiện được chia theo giai đoạn nhằm đảm bảo hoàn thành đúng phạm vi:")
    add_bullets(
        doc,
        [
            "Tuần 1: phân tích yêu cầu, thiết kế CSDL, xây dựng UML (use-case, class, sequence).",
            "Tuần 2: triển khai chức năng nền tảng (tour, user, admin CRUD), giao diện chính và kết nối DB.",
            "Tuần 3: tích hợp VNPay, QR check-in, API bên thứ ba (Amadeus/News), chatbot AI.",
            "Tuần 4: hoàn thiện UI, sửa lỗi, deploy thử nghiệm, viết báo cáo và chuẩn bị demo.",
        ],
    )

    doc.add_heading("Đảm bảo làm việc nhóm", level=2)
    add_bullets(
        doc,
        [
            "Phân công theo module: User (backend + frontend) và Admin (backend + frontend).",
            "Sử dụng Git/GitHub để quản lý mã nguồn, hạn chế commit trực tiếp lên nhánh chính.",
            "Thống nhất quy ước đặt tên và cấu trúc code; cấu hình nhạy cảm được đưa vào .gitignore.",
        ],
    )

    doc.add_heading("Các vấn đề về đạo đức và làm việc chuyên nghiệp", level=2)
    add_bullets(
        doc,
        [
            "Không công khai mật khẩu, API key, secret; ưu tiên dùng biến môi trường khi deploy.",
            "Mã hóa mật khẩu bằng BCrypt; phân quyền truy cập admin.",
            "Tôn trọng dữ liệu người dùng: chỉ thu thập thông tin cần thiết cho đặt tour; không sử dụng sai mục đích.",
        ],
    )

    doc.add_heading("Tác động xã hội", level=2)
    add_bullets(
        doc,
        [
            "Tích cực: giúp người dùng tra cứu minh bạch, giảm thời gian đặt tour; hỗ trợ doanh nghiệp số hóa vận hành.",
            "Lưu ý: cần đảm bảo bảo mật dữ liệu và hỗ trợ trải nghiệm cho người dùng ít tiếp cận công nghệ.",
        ],
    )

    doc.add_heading("Kế hoạch cho kiến thức mới và chiến lược học tập", level=2)
    add_bullets(
        doc,
        [
            "Mở rộng bảo mật: JWT, refresh token, audit log; chính sách CSRF/Rate limit cho API.",
            "Tối ưu hiệu năng: caching (Redis), tối ưu truy vấn, bổ sung index.",
            "Hoàn thiện kiểm thử: unit test/integration test cho booking và payment callback.",
            "Nâng cấp chatbot: RAG với vector DB, cơ chế kiểm soát nội dung (guardrails).",
        ],
    )

    doc.add_heading("Kết luận", level=1)
    doc.add_paragraph(
        "Đồ án đã xây dựng thành công hệ thống ZakiBooking với luồng nghiệp vụ đặt tour tương đối đầy đủ: tìm kiếm - xem chi tiết - đặt tour - "
        "áp mã giảm giá - thanh toán VNPay (sandbox) - theo dõi booking - check-in QR. Hệ thống có phân hệ Admin để quản trị dữ liệu tour và vận hành đơn đặt."
    )
    doc.add_paragraph(
        "Bên cạnh các chức năng nền tảng, hệ thống tích hợp các thành phần nâng cao (OAuth2, chatbot AI, Amadeus/News API) giúp sản phẩm gần với "
        "yêu cầu triển khai thực tế. Kiến trúc MVC phân tầng, MySQL + Flyway, Spring Security giúp hệ thống có cấu trúc rõ ràng, dễ bảo trì và dễ mở rộng."
    )
    doc.add_paragraph(
        "Hướng phát triển tiếp theo gồm: tối ưu hiệu năng và truy vấn khi dữ liệu lớn; hoàn thiện các quy trình vận hành (hủy/hoàn tiền); tách frontend SPA; "
        "tăng cường kiểm thử tự động và nâng cấp chatbot theo mô hình RAG hoàn chỉnh."
    )

    doc.add_heading("Tài liệu tham khảo", level=1)
    add_bullets(
        doc,
        [
            "Spring Boot Documentation: https://docs.spring.io/spring-boot/",
            "Spring Security Reference: https://docs.spring.io/spring-security/",
            "Thymeleaf Documentation: https://www.thymeleaf.org/documentation.html",
            "Flyway Documentation: https://documentation.red-gate.com/fd",
            "MySQL Manual: https://dev.mysql.com/doc/",
            "VNPay Sandbox: https://sandbox.vnpayment.vn/",
            "Amadeus for Developers: https://developers.amadeus.com/",
            "Bootstrap 5: https://getbootstrap.com/docs/5.3/",
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    src = Path(r"C:\Users\minhd\Downloads\Báo cáo đồ án.docx")
    out = Path(r"C:\Users\minhd\Downloads\Báo cáo đồ án - cải tiến.docx")
    upgrade_report(src, out)
    print(f"Wrote: {out}")
