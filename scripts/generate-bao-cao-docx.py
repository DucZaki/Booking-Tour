# -*- coding: utf-8 -*-
"""Generate improved ZakiBooking report docx from updated content."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(r"c:\Users\minhd\Downloads\Báo cáo đồ án - CẢI TIẾN.docx")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(13)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(13)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    return table


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Cover
    for line, size, bold in [
        ("ĐẠI HỌC PHENIKAA", 14, True),
        ("KHOA HỆ THỐNG THÔNG TIN", 14, True),
        ("", 12, False),
        ("BÁO CÁO ĐỒ ÁN LIÊN NGÀNH", 14, True),
        ("Website Booking Tour Du Lịch - ZakiBooking", 13, True),
        ("", 12, False),
        ("Danh sách thành viên:", 13, False),
        ("Nguyễn Minh Đức – 23010634", 13, False),
        ("Giảng viên hướng dẫn: TS. Đặng Thị Thúy An", 13, False),
        ("Thời gian thực hiện: 09/04/2026 – 17/06/2026 (10 tuần)", 13, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)

    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(
        doc,
        "Trong quá trình thực hiện đồ án với đề tài “Xây dựng Website đặt tour du lịch trực tuyến ZakiBooking”, "
        "em đã nhận được sự hỗ trợ, hướng dẫn và động viên từ nhiều phía.",
    )
    add_para(
        doc,
        "Em xin bày tỏ lòng biết ơn sâu sắc tới TS. Đặng Thị Thúy An – giảng viên hướng dẫn, "
        "đã tận tình định hướng, góp ý và theo sát tiến độ thực hiện đồ án.",
    )
    add_para(
        doc,
        "Em cũng xin cảm ơn các thầy cô Khoa Hệ thống thông tin đã truyền đạt kiến thức nền tảng; "
        "cùng gia đình và bạn bè đã hỗ trợ, tạo điều kiện để em hoàn thành đồ án.",
    )
    add_para(
        doc,
        "Do hạn chế về thời gian và kinh nghiệm, báo cáo không tránh khỏi thiếu sót. "
        "Em rất mong nhận được ý kiến góp ý từ thầy cô.",
    )

    add_heading(doc, "1. Giới thiệu", 1)
    add_heading(doc, "1.3. Giải pháp đề xuất (ZakiBooking) — Bổ sung", 2)
    add_para(doc, "Ngoài các chức năng cơ bản, phiên bản hiện tại của ZakiBooking tích hợp thêm:")
    add_bullets(
        doc,
        [
            "Chuyến đi gần bạn: gợi ý tour theo vị trí GPS (Geolocation + Haversine).",
            "Báo giá vé theo điểm đón × ngày khởi hành (Amadeus API, bảng ngay_khoi_hanh_diem_don).",
            "Mã giảm giá linh hoạt: PERCENT và FIXED, áp dụng toàn hệ thống hoặc theo tour.",
            "QR Check-in: sinh mã UUID sau thanh toán VNPay, xác thực tại điểm tập trung.",
            "Chatbot AI (Zaki AI): RAG-lite, fallback Groq → OpenRouter → Gemini.",
            "Thanh toán VNPay sandbox, email xác nhận, OAuth2 Google/Facebook, News API.",
        ],
    )

    add_heading(doc, "2. Thiết kế và triển khai — Phần bổ sung", 1)

    add_heading(doc, "2.1. Yêu cầu chức năng mới", 2)
    add_heading(doc, "2.1.8. Gợi ý chuyến đi gần bạn", 3)
    add_bullets(
        doc,
        [
            "Lấy tọa độ người dùng qua Geolocation API hoặc chọn thủ công Hà Nội / HCM / Đà Nẵng.",
            "API GET /api/tour/nearby lọc tour theo điểm đón trong bán kính (mặc định 100 km).",
            "Hiển thị khoảng cách ước lượng trên trang chủ.",
        ],
    )

    add_heading(doc, "2.1.9. Báo giá vé theo điểm đón", 3)
    add_bullets(
        doc,
        [
            "Admin đồng bộ giá vé Amadeus theo từng điểm đón trên màn hình ngày khởi hành.",
            "Khách chọn điểm đón và ngày KH → AJAX cập nhật giá realtime trên form đặt tour.",
            "API GET /api/tour/{id}/flight-quote?nkhId=&diemDonId=.",
        ],
    )

    add_heading(doc, "2.1.10. Mã giảm giá, QR Check-in, Chatbot AI", 3)
    add_bullets(
        doc,
        [
            "Mã giảm giá PERCENT/FIXED; validate qua PromoApiController.",
            "Sau thanh toán PAID: sinh ma_check_in (UUID), hiển thị QR (ZXing).",
            "Chatbot inject dữ liệu tour/mã giảm giá từ DB; không bịa thông tin.",
        ],
    )

    add_heading(doc, "2.5.1. Bảng Use-case cập nhật", 2)
    add_table(
        doc,
        ["Mã UC", "Tên kịch bản", "Mô tả ngắn"],
        [
            ["UC01", "Xem trang chủ", "Tour nổi bật, Chuyến đi gần bạn, tin tức"],
            ["UC02", "Tìm kiếm/lọc tour", "Lọc điểm đến, giá, ngày KH"],
            ["UC05", "Đặt tour", "Chọn ngày KH, điểm đón, mã giảm giá"],
            ["UC06", "Thanh toán VNPay", "Redirect sandbox, callback/IPN"],
            ["UC07", "Quản lý đơn đặt", "Lịch sử booking, QR check-in"],
            ["UC10", "Chatbot AI", "Tư vấn tour theo dữ liệu hệ thống"],
            ["UC11", "Liên hệ", "Form liên hệ/tư vấn"],
            ["UC-A03", "Quản lý ngày KH", "Điểm đón × giá vé, sync Amadeus"],
            ["UC-A05", "Mã giảm giá", "CRUD mã PERCENT/FIXED"],
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "2.5.3. Mô hình lớp — Entity bổ sung", 2)
    add_table(
        doc,
        ["Entity", "Bảng DB", "Vai trò"],
        [
            ["DiemDon", "diem_don", "Điểm đón (HN, HCM, ĐN…)"],
            ["NgayKhoiHanh", "ngay_khoi_hanh", "Ngày đi/về của tour"],
            ["NgayKhoiHanhDiemDon", "ngay_khoi_hanh_diem_don", "Giá vé theo điểm đón × ngày KH"],
            ["MaGiamGia", "ma_giam_gia", "Mã khuyến mãi PERCENT/FIXED"],
            ["DatCho", "dat_cho", "Đơn đặt, thanh toán, ma_check_in"],
        ],
    )
    add_para(doc, "Chèn hình: docs/diagrams/class-diagram-tour.png, class-diagram-user.png", bold=True)

    add_heading(doc, "2.5.4. Biểu đồ tuần tự mới", 2)
    add_bullets(
        doc,
        [
            "Hình 2.X — Báo giá vé điểm đón: seq-dat-tour-quote.png",
            "Hình 2.X — Thanh toán VNPay: seq-vnpay-payment.png",
            "Hình 2.X — Admin sync Amadeus: seq-admin-amadeus.png",
        ],
    )

    add_heading(doc, "2.6. Triển khai kỹ thuật và tích hợp", 1)
    add_table(
        doc,
        ["Thành phần", "Phiên bản", "Vai trò"],
        [
            ["Java", "17", "Backend"],
            ["Spring Boot", "3.5.6", "Framework"],
            ["MySQL + Flyway", "V1–V7", "CSDL, migration"],
            ["VNPay", "Sandbox", "Thanh toán"],
            ["Amadeus API", "OAuth2", "Báo giá vé"],
            ["Groq/OpenRouter/Gemini", "—", "Chatbot AI"],
            ["Bootstrap 5", "—", "UI responsive, dark mode"],
            ["ZXing", "3.5.3", "QR check-in"],
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "2.6.4. Tính năng nâng cao (tóm tắt)", 2)
    sections = [
        (
            "A. Chuyến đi gần bạn",
            "TourService + GeoUtils (Haversine) + nearby-tours.js. Fallback chọn thành phố khi không có GPS.",
        ),
        (
            "B. Giá vé điểm đón × ngày KH",
            "Bảng ngay_khoi_hanh_diem_don; admin sync Amadeus; user AJAX flight-quote.",
        ),
        (
            "C. Mã giảm giá",
            "PERCENT/FIXED; ap_dung_tat_ca hoặc ma_giam_gia_tour; validate trước VNPay.",
        ),
        (
            "D. QR Check-in",
            "UUID sau PAID; CheckInService xác thực một lần; hiển thị trên bookings.html.",
        ),
        (
            "E. Chatbot AI",
            "ChatService RAG-lite; prompt ràng buộc; fallback provider chain.",
        ),
    ]
    for title, body in sections:
        add_para(doc, title, bold=True)
        add_para(doc, body)

    add_heading(doc, "2.7. Kiểm thử hệ thống", 1)
    add_table(
        doc,
        ["Loại kiểm thử", "Số ca", "Kết quả"],
        [
            ["Chức năng User + Admin", "19", "~100%"],
            ["Tích hợp (VNPay, Amadeus, AI…)", "9", "~100%"],
            ["UI/UX", "6", "Đạt"],
            ["Bảo mật", "5", "Đạt"],
            ["Unit test JUnit", "2", "100%"],
        ],
    )
    add_para(
        doc,
        "Các tính năng mới (nearby, flight-quote, QR, chatbot, promo) đã được kiểm thử thủ công trên localhost:8080.",
    )

    add_heading(doc, "3. Một số thành phần khác — Cập nhật (đồ án cá nhân)", 1)
    add_heading(doc, "3.1. Kế hoạch dự án (10 tuần)", 2)
    weeks = [
        ("1", "09/04–15/04", "Phân tích yêu cầu", "Bảng yêu cầu chức năng"),
        ("2", "16/04–22/04", "Thiết kế Use-case, ERD", "Sơ đồ UML"),
        ("3", "23/04–29/04", "Class/Sequence diagram", "Flyway V1"),
        ("4", "30/04–06/05", "Spring Boot, Entity", "CSDL MySQL"),
        ("5", "07/05–13/05", "Giao diện User", "Xem/tìm tour"),
        ("6", "14/05–20/05", "Security, Admin CRUD", "Dashboard admin"),
        ("7", "21/05–27/05", "Đặt tour, mã giảm giá", "Booking end-to-end"),
        ("8", "28/05–03/06", "VNPay, email, QR", "Thanh toán online"),
        ("9", "04/06–10/06", "Amadeus, AI, Nearby", "Tính năng 2.6.4"),
        ("10", "11/06–17/06", "Kiểm thử, báo cáo", "Deploy Railway"),
    ]
    add_table(doc, ["Tuần", "Thời gian", "Giai đoạn", "Sản phẩm"], weeks)
    doc.add_paragraph()

    add_heading(doc, "3.2. Quy trình phát triển cá nhân", 2)
    add_para(
        doc,
        "Em Nguyễn Minh Đức tự quản lý Git/GitHub, phát triển theo nhánh feature, "
        "self-review trước khi merge vào main. Không commit API key hay mật khẩu DB.",
    )

    add_heading(doc, "4. Kết luận — Cập nhật", 1)
    add_para(
        doc,
        "Đồ án đã xây dựng thành công hệ thống ZakiBooking với đầy đủ luồng nghiệp vụ đặt tour trực tuyến "
        "và các tính năng nâng cao: Chuyến đi gần bạn, báo giá vé theo điểm đón, mã giảm giá linh hoạt, "
        "QR check-in và chatbot AI. Kiến trúc MVC phân tầng kết hợp Spring Security, MySQL và Flyway "
        "đảm bảo khả năng mở rộng và bảo trì.",
    )
    add_para(
        doc,
        "Hạn chế: chưa có E2E test tự động; một số tích hợp phụ thuộc API key bên ngoài. "
        "Hướng phát triển: CI/CD, tối ưu hiệu năng, JWT, vector DB cho chatbot.",
    )

    add_heading(doc, "5. Tài liệu tham khảo", 1)
    refs = [
        "Spring Boot Documentation — https://docs.spring.io/spring-boot/",
        "Spring Security — https://docs.spring.io/spring-security/",
        "VNPay Sandbox — https://sandbox.vnpayment.vn/",
        "Amadeus for Developers — https://developers.amadeus.com/",
        "Flyway — https://documentation.red-gate.com/fd",
        "Bootstrap 5 — https://getbootstrap.com/docs/5.3/",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"{i}. {ref}")

    add_heading(doc, "PHỤ LỤC — Hướng dẫn cập nhật từ báo cáo cũ", 1)
    add_para(doc, "Giữ nguyên từ báo cáo cũ:", bold=True)
    add_bullets(
        doc,
        [
            "Đặc tả Use-case chi tiết (UC1.1–UC6.2) — vẫn hợp lệ.",
            "Biểu đồ tuần tự cũ (đăng nhập, CRUD tour, đặt tour…) — giữ làm nền.",
            "Ảnh chụp màn hình UI đã có trong Word.",
            "Lời cảm ơn, mục 1.1, 1.2 — chỉ sửa “nhóm/chúng em” → “em”.",
        ],
    )
    add_para(doc, "Thêm/chỉnh trong Word cũ:", bold=True)
    add_bullets(
        doc,
        [
            "Mục 1.3: liệt kê 5 tính năng nâng cao (mục 2.6.4).",
            "Mục 2.1: thêm 2.1.8–2.1.10 (nearby, flight-quote, QR/AI/promo).",
            "Mục 2.5.3: bổ sung DiemDon, NgayKhoiHanhDiemDon; chèn class diagram PNG mới.",
            "Mục 2.5.4: thêm 3 sequence diagram mới (VNPay, flight-quote, Amadeus).",
            "Thêm mục 2.6 và 2.7 (copy từ file này hoặc docs/BAO-CAO-DO-AN.md).",
            "Mục 3: đổi làm việc nhóm → cá nhân; kế hoạch 10 tuần (bảng trên).",
            "Mục 4 Kết luận: nhắc tính năng mới.",
            "Use-case diagram: thay bằng usecase-uml-khach-hang.png, usecase-uml-admin.png.",
            "Chụp thêm: section Chuyến đi gần bạn, chatbot, QR, màn admin ngày KH/promo.",
        ],
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
