import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _strip_md_inline(s: str) -> str:
    # Keep it simple: remove common inline markers.
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = s.replace("\\*", "*")
    return s


def _is_table_line(line: str) -> bool:
    line = line.rstrip("\n")
    return line.strip().startswith("|") and line.strip().endswith("|")


def md_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    in_code_block = False
    pending_table: list[str] = []

    def flush_table():
        nonlocal pending_table
        if not pending_table:
            return
        # Write table as plain paragraphs to avoid complex markdown->table parsing.
        doc.add_paragraph("BANG (trich tu Markdown):")
        for l in pending_table:
            doc.add_paragraph(_strip_md_inline(l))
        pending_table = []

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            flush_table()
            in_code_block = not in_code_block
            p = doc.add_paragraph(_strip_md_inline(line))
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in doc.styles else doc.styles["Normal"]
            continue

        if in_code_block:
            # Keep code formatting minimal
            p = doc.add_paragraph(line)
            p.style = doc.styles["No Spacing"] if "No Spacing" in doc.styles else doc.styles["Normal"]
            continue

        if _is_table_line(line) or (pending_table and line.strip().startswith("|")):
            pending_table.append(line)
            continue
        else:
            flush_table()

        if not line.strip():
            doc.add_paragraph("")
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(_strip_md_inline(line[2:].strip()), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("## "):
            doc.add_heading(_strip_md_inline(line[3:].strip()), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(_strip_md_inline(line[4:].strip()), level=3)
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph("" + "_" * 50)
            continue

        # Bullets
        if line.lstrip().startswith("- "):
            content = _strip_md_inline(line.lstrip()[2:].strip())
            doc.add_paragraph(content, style="List Bullet")
            continue

        # Numbered (simple)
        if re.match(r"^\d+\.\s+", line.strip()):
            content = re.sub(r"^\d+\.\s+", "", line.strip())
            doc.add_paragraph(_strip_md_inline(content), style="List Number")
            continue

        # Normal paragraph
        doc.add_paragraph(_strip_md_inline(line))

    flush_table()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "BAO-CAO-DO-AN-CHI-TIET.md"
    out = root / "docs" / "BAO-CAO-DO-AN-CHI-TIET.docx"
    md_to_docx(md, out)
    print(f"Wrote: {out}")
